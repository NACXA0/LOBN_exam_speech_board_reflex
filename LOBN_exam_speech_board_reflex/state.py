# Application state management

from __future__ import annotations

import os
import json
import time
import base64, json
from typing import List, Optional

import reflex as rx

from LOBN_exam_speech_board_reflex.data.question_bank import (
    QuestionBank,
    Question,
    load_question_bank,
    save_question_bank,
    delete_question_bank,
    rename_question_bank,
    get_all_bank_files,
    ensure_directories,
)
from LOBN_exam_speech_board_reflex.data.importers import (
    import_text_content,
    import_json_content,
    import_docx_content,
    import_doc_content,
    get_import_files,
    clean_import_file,
    image_to_base64,
)
import urllib.parse


class AppState(rx.State):
    """Base state shared across all pages."""

    # Current question bank
    current_bank: dict = {}
    current_bank_filename: str = ""

    # Current question index
    current_index: int = 0

    # UI state
    selected_option: int = -1
    show_explanation: bool = False
    right_panel_collapsed: bool = False

    # 页面布局模式: "classic" = 旧双栏布局, "speech" = 新讲题布局(默认)
    workspace_layout: str = "speech"
    # 右侧面板Tab: "explanation" = 全部展示解析, "half" = 解析演讲稿各一半, "speech" = 全部展示演讲稿
    right_panel_tab: str = "half"

    # Answer history for statistics
    answer_history: dict = {}  # {question_id: {"selected": option_idx, "correct": bool}}

    # Bank list for wellcome Tab
    bank_list: list[dict] = []

    # Upload state
    upload_progress: int = 0
    upload_status: str = ""
    preview_bank: dict = {}

    # File upload state
    uploaded_file: Optional[dict] = None

    # Import files list
    import_files: list[dict] = []

    # Tab2 操作反馈状态（需定义在AppState，因为操作方法在此类中）
    action_status: str = ""

    # 白板背景设置
    whiteboard_bg_color: str = "#ffffff"
    whiteboard_watermark: str = ""
    whiteboard_watermark_image: str = ""  # 水印图片(base64或路径)
    whiteboard_watermark_opacity: int = 50  # 水印图片透明度(0-100)，默认50%
    whiteboard_watermark_text_opacity: int = 50  # 水印文字透明度(0-100)，默认50%

    def _load_whiteboard_config(self):
        """Load whiteboard background config from file."""
        try:
            from LOBN_exam_speech_board_reflex.data.question_bank import DATA_DIR
            fpath = os.path.join(DATA_DIR, "whiteboard_config.json")
            if os.path.exists(fpath):
                with open(fpath, "r", encoding="utf-8") as f:
                    config = json.load(f)
                self.whiteboard_bg_color = config.get("bg_color", "#ffffff")
                self.whiteboard_watermark = config.get("watermark", "")
                self.whiteboard_watermark_image = config.get("watermark_image", "")
                self.whiteboard_watermark_opacity = config.get("watermark_opacity", 50)
                self.whiteboard_watermark_text_opacity = config.get("watermark_text_opacity", 50)
        except Exception:
            pass

    def _save_whiteboard_config(self):
        """Save whiteboard background config to file."""
        try:
            from LOBN_exam_speech_board_reflex.data.question_bank import DATA_DIR
            ensure_directories()
            fpath = os.path.join(DATA_DIR, "whiteboard_config.json")
            config = {
                "bg_color": self.whiteboard_bg_color,
                "watermark": self.whiteboard_watermark,
                "watermark_image": self.whiteboard_watermark_image,
                "watermark_opacity": self.whiteboard_watermark_opacity,
                "watermark_text_opacity": self.whiteboard_watermark_text_opacity,
            }
            with open(fpath, "w", encoding="utf-8") as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存白板配置失败：{e}")

    @rx.event
    def load_whiteboard_settings(self):
        """Reload whiteboard settings from file."""
        self._load_whiteboard_config()

    @rx.event
    def set_whiteboard_bg_color(self, color: str):
        """Set whiteboard background color."""
        self.whiteboard_bg_color = color
        self._save_whiteboard_config()

    @rx.event
    def set_whiteboard_watermark(self, text: str):
        """Set whiteboard watermark text."""
        self.whiteboard_watermark = text
        self._save_whiteboard_config()

    @rx.event
    def set_whiteboard_watermark_opacity(self, value: list):
        """Set whiteboard watermark image opacity (0-100)."""
        self.whiteboard_watermark_opacity = int(value[0]) if isinstance(value, list) else int(value)
        self._save_whiteboard_config()

    @rx.event
    def set_whiteboard_watermark_text_opacity(self, value: list):
        """Set whiteboard watermark text opacity (0-100)."""
        self.whiteboard_watermark_text_opacity = int(value[0]) if isinstance(value, list) else int(value)
        self._save_whiteboard_config()

    @rx.var
    def watermark_opacity_css(self) -> str:
        """Get watermark opacity as CSS value (0-1)."""
        return f"{self.whiteboard_watermark_opacity / 100}"

    @rx.var
    def watermark_text_color(self) -> str:
        """Get watermark text color with dynamic opacity."""
        opacity = self.whiteboard_watermark_text_opacity / 100
        return f"rgba(0,0,0,{opacity})"

    @rx.var
    def whiteboard_bg_style(self) -> dict:
        """Get whiteboard background style dict."""
        style: dict = {}
        if self.whiteboard_bg_color:
            style["background_color"] = self.whiteboard_bg_color
        return style

    @rx.event(background=True)
    async def initialize(self):
        """Load available question banks on app start."""
        raw_banks = get_all_bank_files()
        formatted = []
        for bank in raw_banks:
            from datetime import datetime
            dt = datetime.fromtimestamp(bank.get("modified", 0))
            formatted.append({
                "filename": bank["filename"],
                "name": bank.get("name", "未命名题库"),
                "total_questions": bank.get("total_questions", 0),
                "modified_str": dt.strftime("%Y-%m-%d %H:%M"),
            })
        async with self:
            self.bank_list = formatted
            self._load_whiteboard_config()

    @rx.event
    def load_bank(self, filename: str):
        """Load a question bank and start the quiz."""
        bank = load_question_bank(filename)
        if bank:
            self.current_bank = bank.to_dict()
            self.current_bank_filename = filename
            self.current_index = 0
            self.selected_option = -1
            self.show_explanation = False
        else:
            raise Exception(f"Failed to load bank: {filename}")

    @rx.event
    def next_question(self):
        """Go to next question."""
        total = self.current_bank.get("total_questions", 0)
        if self.current_index < total - 1:
            self.current_index += 1
            self.selected_option = -1
            self.show_explanation = False
            # Reset answer history for current question
            questions = self.current_bank.get("questions", [])
            if 0 <= self.current_index < len(questions):
                question_id = questions[self.current_index].get("id", self.current_index)
                if question_id in self.answer_history:
                    del self.answer_history[question_id]

    @rx.event
    def prev_question(self):
        """Go to previous question."""
        if self.current_index > 0:
            self.current_index -= 1
            self.selected_option = -1
            self.show_explanation = False
            # Reset answer history for current question
            questions = self.current_bank.get("questions", [])
            if 0 <= self.current_index < len(questions):
                question_id = questions[self.current_index].get("id", self.current_index)
                if question_id in self.answer_history:
                    del self.answer_history[question_id]

    @rx.event
    def go_to_first_question(self):
        """Go to the first question."""
        self.current_index = 0
        self.selected_option = -1
        self.show_explanation = False
        # Reset answer history for current question
        questions = self.current_bank.get("questions", [])
        if 0 <= self.current_index < len(questions):
            question_id = questions[self.current_index].get("id", self.current_index)
            if question_id in self.answer_history:
                del self.answer_history[question_id]

    @rx.event
    def go_to_last_question(self):
        """Go to the last question."""
        total = self.current_bank.get("total_questions", 0)
        if total > 0:
            self.current_index = total - 1
            self.selected_option = -1
            self.show_explanation = False
            # Reset answer history for current question
            questions = self.current_bank.get("questions", [])
            if 0 <= self.current_index < len(questions):
                question_id = questions[self.current_index].get("id", self.current_index)
                if question_id in self.answer_history:
                    del self.answer_history[question_id]

    @rx.event
    def select_option(self, index: int):
        """Handle option selection."""
        self.selected_option = index
        self.show_explanation = True
        
        # Record answer for statistics
        questions = self.current_bank.get("questions", [])
        if 0 <= self.current_index < len(questions):
            question_id = questions[self.current_index].get("id", self.current_index)
            self.answer_history[question_id] = {
                "selected": index,
                "correct": False,  # Will be updated after checking
            }

    @rx.event
    def check_answer(self):
        """Check if the selected answer is correct."""
        questions = self.current_bank.get("questions", [])
        if 0 <= self.current_index < len(questions):
            question = questions[self.current_index]
            correct_answer = question.get("answer", "")
            selected_option = self.selected_option
            
            # Find which option matches the correct answer
            options = question.get("options", [])
            for i, opt in enumerate(options):
                if opt == correct_answer or opt.lower() == correct_answer.lower():
                    self.answer_history[question.get("id", self.current_index)]["correct"] = (i == selected_option)
                    break
        
        # Reset selection after checking
        self.selected_option = -1

    @rx.event
    def reset_quiz(self):
        """Reset current quiz state."""
        self.current_index = 0
        self.selected_option = -1
        self.show_explanation = False
        self.answer_history = {}

    @rx.event
    def toggle_right_panel(self):
        """Toggle right panel collapsed state."""
        self.right_panel_collapsed = not self.right_panel_collapsed

    @rx.event
    def set_workspace_layout(self, layout: str):
        """Set workspace layout mode."""
        self.workspace_layout = layout

    @rx.event
    def set_right_panel_tab(self, tab: str):
        """Set right panel tab state."""
        self.right_panel_tab = tab

    @rx.event
    def set_whiteboard_watermark_image(self, image: str):
        """Set whiteboard watermark image."""
        self.whiteboard_watermark_image = image
        self._save_whiteboard_config()

    @rx.event
    async def handle_watermark_image_upload(self, files: list[rx.UploadFile]):
        """Handle watermark image upload."""
        if not files:
            return
        file = files[0]
        try:
            content = await file.read()
            import mimetypes
            mime_type, _ = mimetypes.guess_type(file.filename or "")
            if mime_type is None:
                mime_type = "image/png"
            b64 = base64.b64encode(content).decode("utf-8")
            data_uri = f"data:{mime_type};base64,{b64}"
            self.whiteboard_watermark_image = data_uri
            self._save_whiteboard_config()
        except Exception as e:
            print(f"上传水印图片失败：{e}")

    @rx.var
    def current_question(self) -> dict:
        """Get current question data for rendering."""
        questions = self.current_bank.get("questions", [])
        if 0 <= self.current_index < len(questions):
            return questions[self.current_index]
        return {}

    @rx.var
    def current_bank_name(self) -> str:
        """Get current bank name for rendering."""
        return self.current_bank.get("name", "未命名题库")

    @rx.var
    def total_questions(self) -> int:
        """Get total question count for rendering."""
        return self.current_bank.get("total_questions", 0)

    @rx.var
    def current_question_text(self) -> str:
        """Get current question text."""
        return self.current_question.get("question", "")

    @rx.var
    def current_question_images(self) -> list[str]:
        """Get current question images."""
        return self.current_question.get("images", [])

    @rx.var
    def current_question_images_list(self) -> list[dict]:
        """Get all images for current question as list of dicts for rx.foreach."""
        return [{"src": img} for img in self.current_question.get("images", [])]

    @rx.var
    def current_question_options(self) -> list[str]:
        """Get current question options."""
        return self.current_question.get("options", [])

    @rx.var
    def current_question_answer(self) -> str:
        """Get current question answer (uppercase and stripped)."""
        return self.current_question.get("answer", "").strip().upper()

    @rx.var
    def current_question_explanation(self) -> str:
        """Get current question explanation."""
        return self.current_question.get("explanation", "")

    @rx.var
    def current_question_speech(self) -> str:
        """Get current question speech script."""
        return self.current_question.get("speech", "")

    @rx.var
    def current_question_id(self) -> str:
        """Get current question ID or index."""
        return str(self.current_question.get("id", self.current_index + 1))

    @rx.var
    def current_question_type(self) -> str:
        """Get current question type (single / multiple / judge)."""
        q_type = self.current_question.get("type", "")
        if q_type:
            return q_type
        from LOBN_exam_speech_board_reflex.data.question_bank import _detect_question_type
        options = self.current_question.get("options", [])
        answer = self.current_question.get("answer", "")
        return _detect_question_type(options, answer)

    @rx.var
    def option_letters(self) -> list[str]:
        """Get list of option letters (A, B, C, D...) for current question."""
        num_options = len(self.current_question.get("options", []))
        return [chr(65 + i) for i in range(num_options)]  # A, B, C, D...

    @rx.var
    def options_with_letters(self) -> list[dict]:
        """Get options with their corresponding letters."""
        options = self.current_question.get("options", [])
        result = []
        for i, opt in enumerate(options):
            result.append({
                "letter": chr(65 + i),
                "text": opt,
                "index": i
            })
        return result

    @rx.var
    def answer_statistics(self) -> dict[str, str | int | bool | float]:
        """Get answer statistics for current quiz."""
        total = len(self.answer_history)
        correct = sum(1 for v in self.answer_history.values() if v.get("correct", False))
        incorrect = total - correct
        accuracy_value = correct/total*100 if total > 0 else 0
        return {
            "total": total,
            "correct": correct,
            "incorrect": incorrect,
            "accuracy": f"{accuracy_value:.1f}%" if total > 0 else "0%",
            "accuracy_value": accuracy_value,  # Numeric value for comparison
            "is_good": accuracy_value >= 70,  # Boolean for color scheme
        }

    # ---- Upload / Import handlers ----
    @rx.event
    async def handle_text_upload(self, text: str, filename: str = "imported.txt"):
        """Process text-based upload."""
        get_AdminState = self.get_state(AdminState)
        try:
            self.upload_status = "正在解析..."
            bank = import_text_content(text, filename)
            # 提取 description（从文件头或作为单独字段）
            desc = getattr(get_AdminState, 'bank_description', '').strip() if hasattr(get_AdminState, 'bank_description') else ""
            if desc:
                # 更新 preview_bank 中的 description
                if "description" not in bank.to_dict():
                    bank_to_dict = bank.to_dict()
                    bank_to_dict["description"] = desc
                    from LOBN_exam_speech_board_reflex.data.question_bank import QuestionBank
                    bank = QuestionBank.from_dict(bank_to_dict)
            self.preview_bank = bank.to_dict()
            self.upload_status = "解析完成，请确认内容"
        except Exception as e:
            self.upload_status = f"解析失败：{str(e)}"
        finally:
            self.upload_progress = 0

    @rx.event
    def refresh_import_files(self):
        """Refresh the import files list."""
        self.import_files = get_import_files()

    @rx.event
    def process_import_file(self, filename: str):
        """Process a file from the import directory."""
        from LOBN_exam_speech_board_reflex.data.question_bank import IMPORT_DIR
        fpath = os.path.join(IMPORT_DIR, filename)
        try:
            if filename.endswith(".doc"):
                bank = import_doc_content(fpath, filename)
            elif filename.endswith(".docx"):
                bank = import_docx_content(fpath, filename)
            elif filename.endswith(".json"):
                with open(fpath, "r", encoding="utf-8") as f:
                    content = f.read()
                bank = import_json_content(content, filename)
            else:
                with open(fpath, "r", encoding="utf-8") as f:
                    content = f.read()
                bank = import_text_content(content, filename)

            bank.filename = bank.name + ".json"
            save_question_bank(bank)
            clean_import_file(filename)
            self.refresh_import_files()
            self.upload_status = f"文件 '{filename}' 处理成功"
        except Exception as e:
            self.upload_status = f"处理失败: {str(e)}"

    @rx.event
    def delete_bank(self, filename: str):
        """Delete a question bank."""
        if delete_question_bank(filename):
            self.bank_list = get_all_bank_files()
            self.action_status = "题库已删除"
        else:
            self.action_status = "删除失败：题库不存在"

    @rx.event
    def rename_bank(self, old_filename: str, new_name: str):
        """Rename a question bank."""
        if rename_question_bank(old_filename, new_name):
            self.bank_list = get_all_bank_files()
            self.action_status = f"已重命名为「{new_name}」"
        else:
            self.action_status = "重命名失败"

    @rx.event
    def save_bank_description(self, filename: str, description: str):
        """Save description for a question bank."""
        try:
            from LOBN_exam_speech_board_reflex.data.question_bank import QUESTION_BANKS_DIR
            fpath = os.path.join(QUESTION_BANKS_DIR, filename)
            if not os.path.exists(fpath):
                self.action_status = "保存描述失败：题库文件不存在"
                return
            
            with open(fpath, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            data["description"] = description
            from LOBN_exam_speech_board_reflex.data.question_bank import save_question_bank
            bank = QuestionBank.from_dict(data)
            save_question_bank(bank)
            
            # 更新 bank_list
            self.bank_list = get_all_bank_files()
            self.action_status = "描述已更新"
        except Exception as e:
            print(f"保存描述失败：{e}")
            self.action_status = f"保存描述失败：{e}"


class AdminState(AppState):
    """State for the admin page. Inherits AppState for shared fields (preview_bank, upload_status, etc.)."""
    upload_tab: str = "text"
    text_content: str = ""
    md_content: str = ""
    json_content: str = ""
    new_bank_name: str = ""
    bank_description: str = ""  # 题库描述
    rename_filename: str = ""
    rename_new_name: str = ""
    show_delete_confirm: str = ""
    preview_mode: bool = False
    show_preview: bool = True  # 在线上传右侧预览面板是否展开
    show_input_panel: bool = True  # 在线上传左侧输入面板是否展开
    current_action_bank_filename: str = ""  # 当前操作的题库文件名
    current_action_type: str = ""  # 当前操作类型：rename, edit_desc, delete
    bank_sort_field: str = "modified"  # 排序字段：modified, name, total_questions
    bank_sort_reverse: bool = True  # 是否倒序

    @staticmethod
    def _format_time(timestamp) -> str:
        """Format timestamp to readable date string."""
        if not timestamp:
            return "未知"
        from datetime import datetime
        dt = datetime.fromtimestamp(timestamp)
        return dt.strftime("%Y-%m-%d %H:%M")

    # ---- Computed vars for proper typing in rx.foreach ----

    @rx.var
    def preview_questions(self) -> list[dict]:
        """Get preview questions list with explicit typing for rx.foreach."""
        questions = self.preview_bank.get("questions", [])
        return questions if isinstance(questions, list) else []

    @rx.var
    def preview_questions_formatted(self) -> list[dict[str, str | bool]]:
        """Get preview questions with formatted display text."""
        questions = self.preview_bank.get("questions", [])
        if not isinstance(questions, list):
            return []
        result = []
        for question in questions:
            q_copy = dict(question)
            q_id = question.get("id", "?")
            question_text = question.get("question", "")
            q_copy["display_brief"] = f"第{q_id}题：{question_text[:50]}"
            q_copy["display_id"] = f"第{q_id}题"
            q_copy["display_question"] = question_text
            # 选项：逐行显示，带字母标号
            options = question.get("options", [])
            q_copy["display_options"] = "\n".join(
                f"{chr(65 + i)}. {opt}" for i, opt in enumerate(options)
            )
            # 选项列表（用于逐行渲染）
            q_copy["display_options_list"] = [
                {"letter": chr(65 + i), "text": opt}
                for i, opt in enumerate(options)
            ]
            # 答案
            answer = question.get("answer", "")
            q_copy["display_answer"] = answer
            q_copy["has_answer"] = bool(answer)
            # 解析
            explanation = question.get("explanation", "")
            q_copy["display_explanation"] = explanation
            q_copy["has_explanation"] = bool(explanation)
            # 选项是否存在
            q_copy["has_options"] = bool(options)
            result.append(q_copy)
        return result

    @rx.var
    def preview_questions_brief(self) -> list[dict]:
        """Get first 5 preview questions with formatted display text."""
        return self.preview_questions_formatted[:5]

    @rx.var
    def preview_total_questions(self) -> int:
        """Get total questions count from preview bank."""
        questions = self.preview_bank.get("questions", [])
        return len(questions) if isinstance(questions, list) else 0

    @rx.var
    def has_preview(self) -> bool:
        """Whether there is a preview bank loaded."""
        return bool(self.preview_bank)

    @rx.var
    def import_files_list(self) -> list[dict]:
        """Get import files list with explicit typing for rx.foreach."""
        return self.import_files if isinstance(self.import_files, list) else []

    @rx.var
    def bank_list_items(self) -> list[dict]:
        """Get bank list with explicit typing for rx.foreach."""
        return self.bank_list if isinstance(self.bank_list, list) else []

    @rx.var
    def bank_list_formatted(self) -> list[dict]:
        """Get bank list with formatted fields for display, sorted by current sort settings."""
        banks = self.bank_list if isinstance(self.bank_list, list) else []
        result = []
        for bank in banks:
            bank_copy = dict(bank)
            description = bank.get("description", "")
            bank_copy["display_description"] = f"{description[:80]}{'...' if len(description) > 80 else ''}" if description else "无描述"
            # 格式化时间
            modified_timestamp = bank.get("modified", 0)
            if modified_timestamp:
                from datetime import datetime
                dt = datetime.fromtimestamp(modified_timestamp)
                bank_copy["display_modified"] = dt.strftime("%Y-%m-%d %H:%M")
            else:
                bank_copy["display_modified"] = "未知"
            # 题目数
            bank_copy["display_total_questions"] = bank.get("total_questions", 0)
            # 文件大小
            size_bytes = bank.get("size", 0)
            if size_bytes >= 1024 * 1024:
                bank_copy["display_size"] = f"{size_bytes / (1024 * 1024):.1f} MB"
            elif size_bytes >= 1024:
                bank_copy["display_size"] = f"{size_bytes / 1024:.1f} KB"
            else:
                bank_copy["display_size"] = f"{size_bytes} B"
            result.append(bank_copy)
        # 排序
        sort_field = self.bank_sort_field
        reverse = self.bank_sort_reverse
        if sort_field == "name":
            result.sort(key=lambda b: b.get("name", "").lower(), reverse=reverse)
        elif sort_field == "total_questions":
            result.sort(key=lambda b: b.get("total_questions", 0), reverse=reverse)
        else:  # modified (default)
            result.sort(key=lambda b: b.get("modified", 0), reverse=reverse)
        return result

    @rx.var
    def bank_sort_indicator(self) -> str:
        """Get sort indicator for current sort field."""
        if self.bank_sort_reverse:
            return "↓"
        return "↑"

    @rx.var
    def import_files_with_size(self) -> list[dict]:
        """Get import files list with formatted size."""
        files = self.import_files if isinstance(self.import_files, list) else []
        result = []
        for f in files:
            file_copy = dict(f)
            size_bytes = f.get("size", 0)
            file_copy["size_kb"] = f"{size_bytes / 1024:.1f} KB"
            result.append(file_copy)
        return result

    @rx.event
    async def handle_inline_file_upload(self, files: list[rx.UploadFile]):
        """Handle file upload from the inline button in online upload tab.
        Parses the file and fills the text content into the input area.
        """
        if not files:
            return

        file = files[0]
        try:
            content = await file.read()
            filename = file.filename or "uploaded.txt"

            # Auto-set bank name from filename (without extension)
            bank_name = os.path.splitext(filename)[0]
            if not self.new_bank_name:
                self.new_bank_name = bank_name

            if filename.endswith(".json"):
                # JSON is standard format: parse directly, fill raw JSON in text area
                text = content.decode("utf-8")
                bank = import_json_content(text, filename)
                self.preview_bank = bank.to_dict()
                self.text_content = text
                self.upload_status = f"文件 '{filename}' 已加载（标准JSON格式）"
            elif filename.endswith(".doc"):
                # Legacy .doc format: save to temp, use import_doc_content
                import tempfile
                with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name
                try:
                    bank = import_doc_content(tmp_path, filename)
                    # Extract text for the input area via win32com or docx
                    try:
                        import win32com.client
                        import pythoncom
                        pythoncom.CoInitialize()
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        doc_obj = word.Documents.Open(os.path.abspath(tmp_path))
                        extracted_text = doc_obj.Content.Text
                        doc_obj.Close(False)
                        word.Quit()
                        pythoncom.CoUninitialize()
                    except Exception:
                        try:
                            import docx as docx_lib
                            doc_obj = docx_lib.Document(tmp_path)
                            extracted_text = "\n".join([para.text for para in doc_obj.paragraphs])
                        except Exception:
                            extracted_text = ""
                finally:
                    os.unlink(tmp_path)
                self.text_content = extracted_text
                if bank and bank.questions:
                    self.preview_bank = bank.to_dict()
                elif extracted_text.strip():
                    bank2 = import_text_content(extracted_text, filename)
                    self.preview_bank = bank2.to_dict()
                else:
                    self.preview_bank = {}
                self.upload_status = f"文件 '{filename}' 已加载"
            elif filename.endswith(".docx"):
                # Extract text from docx, fill into text area, auto-preview
                import tempfile
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name
                try:
                    import docx as docx_lib
                    doc = docx_lib.Document(tmp_path)
                    extracted_text = "\n".join([para.text for para in doc.paragraphs])
                except ImportError:
                    extracted_text = ""
                finally:
                    os.unlink(tmp_path)
                self.text_content = extracted_text
                if extracted_text.strip():
                    bank = import_text_content(extracted_text, filename)
                    self.preview_bank = bank.to_dict()
                else:
                    self.preview_bank = {}
                self.upload_status = f"文件 '{filename}' 已加载"
            else:
                # .txt / .md: fill text area directly, auto-preview
                text = content.decode("utf-8")
                self.text_content = text
                if text.strip():
                    bank = import_text_content(text, filename)
                    self.preview_bank = bank.to_dict()
                else:
                    self.preview_bank = {}
                self.upload_status = f"文件 '{filename}' 已加载"
        except Exception as e:
            self.upload_status = f"文件读取失败：{str(e)}"

    @rx.event
    def set_upload_tab(self, value: str):
        """Set the active upload tab."""
        self.upload_tab = value

    @rx.event
    def set_text_content(self, text: str):
        """Set text content and auto-generate real-time preview."""
        self.text_content = text
        if text.strip():
            try:
                filename = self.new_bank_name + ".txt" if self.new_bank_name else "untitled.txt"
                bank = import_text_content(text, filename)
                self.preview_bank = bank.to_dict()
            except Exception:
                self.preview_bank = {}
        else:
            self.preview_bank = {}

    @rx.event
    def set_rename_new_name(self, name: str):
        """Set rename new name state."""
        self.rename_new_name = name

    @rx.event
    def confirm_preview(self):
        """Confirm and save the previewed bank."""
        if self.preview_bank:
            bank = QuestionBank.from_dict(self.preview_bank)
            if self.new_bank_name:
                bank.name = self.new_bank_name
                bank.filename = f"{self.new_bank_name}.json"
            save_question_bank(bank)
            self.upload_status = "题库保存成功"
            self.preview_bank = {}
            self.text_content = ""
            self.new_bank_name = ""
            self.bank_description = ""
            self.uploaded_file = None
            # 刷新bank list
            self.bank_list = get_all_bank_files()

    @rx.event
    def discard_preview(self):
        """Discard the preview and reset form."""
        self.preview_bank = {}
        self.upload_status = ""
        self.uploaded_file = None

    @rx.event
    def set_bank_sort(self, field: str):
        """Set bank list sort field and toggle order."""
        if self.bank_sort_field == field:
            self.bank_sort_reverse = not self.bank_sort_reverse
        else:
            self.bank_sort_field = field
            self.bank_sort_reverse = (field != "name")  # 名称默认正序，其他默认倒序

    @rx.event
    def set_show_delete_confirm(self, value: str):
        """Set the delete confirm dialog."""
        self.show_delete_confirm = value

    @rx.event
    def open_action_dialog(self, action_type: str, filename: str):
        """Open action dialog for a specific bank."""
        self.current_action_type = action_type
        self.current_action_bank_filename = filename
        # 如果是重命名，预填充当前名称
        if action_type == "rename":
            for bank in self.bank_list:
                if bank.get("filename") == filename:
                    self.rename_new_name = bank.get("name", "")
                    break
        # 如果是编辑描述，预填充当前描述
        if action_type == "edit_desc":
            for bank in self.bank_list:
                if bank.get("filename") == filename:
                    self.bank_description = bank.get("description", "")
                    break
        self.action_status = ""

    @rx.event
    def close_action_dialog(self):
        """Close action dialog."""
        self.current_action_type = ""
        self.current_action_bank_filename = ""

    @rx.event
    def set_bank_description(self, value: str):
        """Set bank description."""
        self.bank_description = value

    @rx.event
    def set_new_bank_name(self, value: str):
        """Set new bank name and update preview if exists."""
        self.new_bank_name = value
        # 如果已有预览，更新题库名称
        if self.preview_bank and value:
            self.preview_bank["name"] = value

    @rx.event
    def toggle_preview(self):
        """Toggle the preview panel visibility."""
        self.show_preview = not self.show_preview

    @rx.event
    def toggle_input_panel(self):
        """Toggle the input panel visibility."""
        self.show_input_panel = not self.show_input_panel

    @rx.event
    def set_show_input_panel(self, value: bool):
        """Set input panel visibility."""
        self.show_input_panel = value

    @rx.event
    def set_show_preview(self, value: bool):
        """Set preview panel visibility."""
        self.show_preview = value

    @rx.event
    def download_sample_json(self):
        """Download a sample JSON file in the system's question bank format."""

        sample_data: dict[str, str | list[dict[str, str | int | list[str]]]] = {
            "name": "示例题库",
            "description": "这是一个示例题库，展示系统支持的JSON格式",
            "questions": [
                {
                    "id": 1,
                    "question": "下列哪个是Python的内置数据类型？",
                    "options": ["Array", "Dictionary", "List", "Both B and C"],
                    "answer": "D",
                    "explanation": "Python的内置数据类型包括字典(dict)和列表(list)。Array不是Python的内置类型。",
                    "speech": "这道题考察Python的基础数据类型。大家注意，Array不是Python内置的，需要import array模块才能用。而字典和列表都是Python最常用的内置数据结构。"
                },
                {
                    "id": 2,
                    "question": "Python中用于定义函数的关键字是？",
                    "options": ["function", "def", "func", "define"],
                    "answer": "B",
                    "explanation": "Python使用def关键字来定义函数。",
                    "speech": "定义函数是Python编程的基础。记住是def，不是function也不是func，这是Python独有的语法风格。"
                },
                {
                    "id": 3,
                    "question": "以下哪个方法用于在列表末尾添加元素？",
                    "options": ["insert()", "append()", "add()", "extend()"],
                    "answer": "B",
                    "explanation": "append()方法用于在列表末尾添加单个元素。",
                    "speech": "append和extend容易混淆。append是添加单个元素到末尾，extend是添加一个序列的所有元素。insert可以在指定位置插入。"
                }
            ]
        }
        json_str = json.dumps(sample_data, ensure_ascii=False, indent=2)
        return rx.download(data=json_str, filename="sample_question_bank.json")

    @rx.event
    def download_bank_json(self, filename: str):
        """Download a question bank's raw JSON file."""
        try:
            from LOBN_exam_speech_board_reflex.data.question_bank import QUESTION_BANKS_DIR
            fpath = os.path.join(QUESTION_BANKS_DIR, filename)
            if not os.path.exists(fpath):
                self.action_status = f"下载失败：文件不存在"
                return
            with open(fpath, "r", encoding="utf-8") as f:
                content = f.read()
            self.action_status = f"正在下载 {filename}"
            return rx.download(data=content, filename=filename)
        except Exception as e:
            self.action_status = f"下载失败：{e}"

    @rx.event
    def open_edit_questions_new_tab(self, filename: str):
        """Open edit questions page in a new tab."""
        encoded = urllib.parse.quote(filename)
        return rx.redirect(f"/edit-questions/{encoded}", is_external=True)

    @rx.event
    def clean_import_file(self, filename: str):
        """Clean an import file."""
        clean_import_file(filename)
        # Refresh import files list
        self.refresh_import_files()


class EditQuestionsState(AppState):
    """State for the edit questions page - supports independent editing per tab."""

    editing_bank_filename: str = ""
    editing_bank_name: str = ""
    editing_bank_questions: list[dict] = []
    edit_status: str = ""
    upload_dialog_open: bool = False
    upload_target_question_index: int = -1
    image_preview_open: bool = False
    image_preview_src: str = ""

    @rx.event
    def open_image_preview(self, src: str):
        """Open image preview dialog with the given image source."""
        self.image_preview_src = src
        self.image_preview_open = True

    @rx.event
    def close_image_preview(self):
        """Close image preview dialog."""
        self.image_preview_open = False
        self.image_preview_src = ""

    @rx.event
    def open_upload_dialog(self, question_index: int):
        """Open the upload dialog for a specific question."""
        self.upload_target_question_index = question_index
        self.upload_dialog_open = True

    @rx.event
    def close_upload_dialog(self):
        """Close the upload dialog."""
        self.upload_dialog_open = False
        self.upload_target_question_index = -1

    @rx.event
    def load_bank_from_url(self):
        """Load bank from route params on page load."""
        bank_filename = self.router.page.params.get("bank", "")
        # URL解码
        if bank_filename:
            bank_filename = urllib.parse.unquote(bank_filename)
        if not bank_filename:
            self.edit_status = "未指定题库文件"
            return
        bank = load_question_bank(bank_filename)
        if bank:
            self.editing_bank_filename = bank_filename
            self.editing_bank_name = bank.name
            self.editing_bank_questions = [q.to_dict() for q in bank.questions]
            self.edit_status = ""
        else:
            self.edit_status = "无法加载题库"

    @rx.event
    def save_edited_question_field(self, question_index: int, field: str, value: str):
        """Save a single field of an edited question (called on_blur)."""
        if 0 <= question_index < len(self.editing_bank_questions):
            self.editing_bank_questions[question_index][field] = value
            self._save_editing_bank()

    @rx.event
    def set_edited_answer(self, question_index: int, letter: str):
        """Set the answer by clicking an option letter (single choice & judge)."""
        if 0 <= question_index < len(self.editing_bank_questions):
            self.editing_bank_questions[question_index]["answer"] = letter
            self._save_editing_bank()

    @rx.event
    def toggle_edited_answer(self, question_index: int, letter: str):
        """Toggle a letter in the answer (multiple choice)."""
        if 0 <= question_index < len(self.editing_bank_questions):
            current = self.editing_bank_questions[question_index].get("answer", "").strip().upper()
            # Split into individual letters
            letters = [c for c in current if c.isalpha()]
            if letter in letters:
                letters.remove(letter)
            else:
                letters.append(letter)
                letters.sort()  # Keep alphabetical order
            self.editing_bank_questions[question_index]["answer"] = "".join(letters)
            self._save_editing_bank()

    @rx.event
    def set_question_type(self, question_index: int, q_type: str):
        """Change the type of a question and adjust data accordingly."""
        if 0 <= question_index < len(self.editing_bank_questions):
            old_type = self.editing_bank_questions[question_index].get("type", "single")
            if q_type == old_type:
                return
            self.editing_bank_questions[question_index]["type"] = q_type
            # 判断题：替换选项为固定 √/×
            if q_type == "judge":
                self.editing_bank_questions[question_index]["options"] = ["√", "×"]
                # 如果当前答案不在合法范围，清空
                ans = self.editing_bank_questions[question_index].get("answer", "").strip().upper()
                if ans not in ("A", "B"):
                    self.editing_bank_questions[question_index]["answer"] = ""
            # 从判断题切回单选/多选：如果选项是固定的 √/×，清空选项
            elif old_type == "judge":
                opts = self.editing_bank_questions[question_index].get("options", [])
                if opts == ["√", "×"] or opts == ["对", "错"]:
                    self.editing_bank_questions[question_index]["options"] = ["", "", "", ""]
                    self.editing_bank_questions[question_index]["answer"] = ""
            # 单选只保留第一个答案字母
            elif q_type == "single":
                ans = self.editing_bank_questions[question_index].get("answer", "").strip().upper()
                if len(ans) > 1:
                    self.editing_bank_questions[question_index]["answer"] = ans[0]
            self._save_editing_bank()

    @rx.event
    def save_edited_option(self, question_index: int, opt_index: int, value: str):
        """Save a single option of an edited question (called on_blur)."""
        if 0 <= question_index < len(self.editing_bank_questions):
            options = self.editing_bank_questions[question_index].get("options", [])
            if 0 <= opt_index < len(options):
                options[opt_index] = value
                self.editing_bank_questions[question_index]["options"] = options
                self._save_editing_bank()

    def _save_editing_bank(self):
        """Persist the editing_bank_questions to disk."""
        if not self.editing_bank_filename:
            return
        try:
            from LOBN_exam_speech_board_reflex.data.question_bank import QUESTION_BANKS_DIR
            fpath = os.path.join(QUESTION_BANKS_DIR, self.editing_bank_filename)
            if not os.path.exists(fpath):
                return
            with open(fpath, "r", encoding="utf-8") as f:
                data = json.load(f)
            data["questions"] = self.editing_bank_questions
            bank = QuestionBank.from_dict(data)
            save_question_bank(bank)
            self.bank_list = get_all_bank_files()
            self.edit_status = "已自动保存"
        except Exception as e:
            self.edit_status = f"自动保存失败：{e}"

    @rx.var
    def editing_questions_formatted(self) -> list[dict[str, str | int | list[dict[str, str | bool | int]]]]:
        """Get editing questions with display formatting."""
        result = []
        for index, question in enumerate(self.editing_bank_questions):
            q_copy = dict(question)
            q_copy["edit_index"] = index
            q_copy["display_id"] = f"第{question.get('id', index+1)}题"
            # 自动检测题型
            question_type = question.get("type", "")
            if not question_type:
                question_type = self._detect_question_type(question)
                q_copy["type"] = question_type
            q_copy["display_type"] = {"single": "单选", "multiple": "多选", "judge": "判断"}.get(question_type, "单选")
            # 预处理答案（避免在 Var 上调用 .strip()）
            q_copy["answer_stripped"] = question.get("answer", "").strip().upper()
            # 选项列表带索引
            options = question.get("options", [])
            answer_str = question.get("answer", "").strip().upper()
            answer_letters = [c for c in answer_str if c.isalpha()]
            q_copy["edit_options"] = [
                {
                    "letter": chr(65 + j),
                    "text": opt,
                    "opt_index": j,
                    "is_answer": chr(65 + j) in answer_letters,
                    "question_index": index,
                    "question_type": question_type,
                }
                for j, opt in enumerate(options)
            ]
            # 图片列表（确保存在）
            q_copy["images"] = question.get("images", [])
            q_copy["has_images"] = bool(question.get("images", []))
            q_copy["edit_images"] = [
                {"src": img, "img_index": k, "question_index": index}
                for k, img in enumerate(question.get("images", []))
            ]
            result.append(q_copy)
        return result

    @staticmethod
    def _detect_question_type(q: dict) -> str:
        """Auto-detect question type from answer and options."""
        answer = q.get("answer", "").strip().upper()
        options = q.get("options", [])
        # 判断题：选项为 √/× 或 对/错
        if options in (["√", "×"], ["对", "错"], ["√", "✗"], ["正确", "错误"]):
            return "judge"
        # 多选：答案包含多个字母
        answer_letters = [c for c in answer if c.isalpha()]
        if len(answer_letters) > 1:
            return "multiple"
        # 默认单选
        return "single"

    @rx.event
    async def handle_upload_dialog_drop(self, files: list[rx.UploadFile]):
        """Handle file drop in the upload dialog."""
        await self.add_image_to_question(self.upload_target_question_index, files)

    @rx.event
    async def add_image_to_question(self, question_index: int, files: list[rx.UploadFile]):
        """Upload image and add to question's images list."""
        if not files or not (0 <= question_index < len(self.editing_bank_questions)):
            self.upload_dialog_open = False
            return
        file = files[0]
        try:
            content = await file.read()
            import mimetypes
            mime_type, _ = mimetypes.guess_type(file.filename or "")
            if mime_type is None:
                mime_type = "image/png"
            b64 = base64.b64encode(content).decode("utf-8")
            data_uri = f"data:{mime_type};base64,{b64}"
            images = self.editing_bank_questions[question_index].get("images", [])
            images.append(data_uri)
            self.editing_bank_questions[question_index]["images"] = images
            self._save_editing_bank()
            self.edit_status = "图片添加成功"
        except Exception as e:
            self.edit_status = f"添加图片失败：{e}"
        finally:
            self.upload_dialog_open = False

    @rx.event
    def remove_image_from_question(self, question_index: int, img_index: int):
        """Remove an image from question's images list."""
        if 0 <= question_index < len(self.editing_bank_questions):
            images = self.editing_bank_questions[question_index].get("images", [])
            if 0 <= img_index < len(images):
                images.pop(img_index)
                self.editing_bank_questions[question_index]["images"] = images
                self._save_editing_bank()

